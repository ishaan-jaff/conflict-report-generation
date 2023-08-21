import json
from docx import Document
from docx.shared import Pt, Inches
from datetime import datetime

class DocxTemplateFiller:
    def __init__(self, template_path, json_path, output_path):
        self.template_path = template_path
        self.json_path = json_path
        self.output_path = output_path

        with open(json_path) as f:
            self.data = json.load(f)

        self.doc = Document(template_path)
        self.bullet_symbol = u'\u2212'  # Minus hyphen symbol
        self.indent_size = Inches(0.5)  # Half inch indentation
        self.font_name = 'Avenir Next LT Pro Light'  # Font name

    def set_run_font(self, run, font_name, font_size):
        run.font.name = font_name
        run.font.size = font_size

    def fill_placeholders(self):
        placeholders = {
            'DD MMM YYYY': (datetime.today().strftime('%d %b %Y'), None),
            'Country Code/ Regional Name / City Name': (self.data.get('LOCATION'), None),
            '(DD, DMS, or MGRS, “Center Point or Quad-Grids”)': (self.data.get('COORDINATES'), None),
            '<INSERT EXECUTIVE SUMMARY>': (self.data.get('EXECUTIVE SUMMARY'), None),
            '<INSERT GPHI EVENTS>': (self.data['theme'].get('General Population & Human Interest'), 'General Population & Human Interest'),
            '<INSERT PCM EVENTS>': (self.data['theme'].get('Public & Civil Movements'), 'Public & Civil Movements'),
            '<INSERT OG EVENTS>': (self.data['theme'].get('Organizations & Governance'), 'Organizations & Governance'),
            '<INSERT MA EVENTS>': (self.data['theme'].get('Military Affairs'), 'Military Affairs'),
            '<INSERT ID EVENTS>': (self.data['theme'].get('Infrastructure & Development'), 'Infrastructure & Development'),
            '<INSERT SC EVENTS>': (self.data['theme'].get('Social & Cultural Insights'), 'Social & Cultural Insights'),
            '<INSERT ED EVENTS>': (self.data['theme'].get('Economic Dynamics'), 'Economic Dynamics'),
            '<INSERT OVERALL SENTIMENT>': (self.data.get('OVERALL SENTIMENT'), None),
            '<INSERT IMPLICATIONS>': (self.data.get('IMPLICATIONS'), None),
            '<INSERT RECOMMENDED ACTIONS>': (self.data.get('RECOMMENDED ACTIONS') and self.data['RECOMMENDED ACTIONS'].split('\n-'), None)
        }
        
        paragraphs_to_delete = []

        for para in self.doc.paragraphs:
            for placeholder, (value, theme) in placeholders.items():
                if value is None:
                    if theme in para.text:
                        paragraphs_to_delete.append(para)
                if placeholder in para.text:
                        if value is None:
                            paragraphs_to_delete.append(para)
                        elif isinstance(value, list):
                            para.clear()  # Clear the existing text
                            for item in value:
                                print(item)
                                if item[0]=='-':
                                    item = item[1:].strip()
                                run = para.add_run(f'{self.bullet_symbol} {item}\n')
                                self.set_run_font(run, self.font_name, Pt(10))  # Set font and size
                                para.paragraph_format.left_indent = self.indent_size  # Set left indentation
                        else:
                            para.text = para.text.replace(placeholder, str(value))
                            for run in para.runs:
                                if placeholder not in run.text:
                                    self.set_run_font(run, self.font_name, run.font.size)
                                    
        for para in paragraphs_to_delete:
            p = para._element
            p.getparent().remove(p)

    def save_output(self):
        self.doc.save(self.output_path)